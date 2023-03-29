from docx.api import Document
import openpyxl
import subprocess
import pandas as pd
import docx
from docx import document
from docx.enum.text import WD_ALIGN_PARAGRAPH

#dbシートより
wb_db = openpyxl.load_workbook('seigikou_db.xlsx')
sheet1 = wb_db['1_event_table']
sheet3 = wb_db['3_kouen_table']
sheet5 = wb_db['5_CPD_table']
sheet7 = wb_db['7_rireki_table']
sheet9 = wb_db['9_member_table']
sheetrei = wb_db['例会']

#dbファイルから値を読み込み
date_value_b = sheet1['b2'].value
date_value = date_value_b.strftime('%Y%m%d')
yearhen = date_value_b.strftime('%Y')
yearstrtoint = int(yearhen)
yeartoreiwa = yearstrtoint - 2018
reiwastr = str(yeartoreiwa)
monthstr = date_value_b.strftime('%m').lstrip("0")
daystr = date_value_b.strftime('%d').lstrip("0")
datevalue_line = '令和' + reiwastr + '年' + monthstr + '月' + daystr + '日'
name_value = sheet1['f2'].value
place_value = sheet1['e2'].value
contents_value = sheet3['e2'].value
start_value = sheet1['c2'].value
end_value = sheet1['d2'].value

#ここからdocx
doc0 = docx.Document("Newcpdformat2.docx")
savedwordname = 'C:\\Users\\niibu\\pythons\\pdfsmerge\\901.docx'
tbl0 = doc0.tables[0]
tbl1 = doc0.tables[1]
row0 = tbl0.rows[0]
row1 = tbl1.rows[1]
row4 = tbl1.rows[4]
row5 = tbl1.rows[5]
row6 = tbl1.rows[6]
row8 = tbl1.rows[8]
row9 = tbl1.rows[9]

p0text ="YO901"+"\r"+datevalue_line
p0=row0.cells[0].add_paragraph(p0text) 
p0.alignment=WD_ALIGN_PARAGRAPH.RIGHT

row4.cells[1].text = datevalue_line
row5.cells[1].text = datevalue_line
row6.cells[1].text = name_value
row8.cells[1].text = place_value
row9.cells[1].text = contents_value

doc0.save(savedwordname)

import gatherdocx2pdf

#ここから過去作成の残骸。ちなみにダブルクォーテーション3つで複数行コメントアウト可能。
"""

#cpd作成用シートより
wb_cpd = openpyxl.load_workbook('cpdsakusei.xlsx')
cpdvalue = wb_cpd['value']
cpdsheet = wb_cpd['CPD']

#CPDファイルに値を書き込み
cpdvalue['b2'].value = datevalue_line
cpdvalue['b5'].value = name_value
cpdvalue['b6'].value = place_value
cpdvalue['b7'].value = contents_value
cpdvalue['b8'].value = start_value
cpdvalue['b9'].value = end_value

#保存する。
wb_cpd.save('cpdsakusei.xlsx')

ukeiretime = sheet0['c3'].value
ukeireday = ukeiretime.strftime('%Y%m%d')
genba = sheet0['C5'].value
zerozeroone = sheet0['d9'].value + '_"' + sheet0['e9'].value + '"'
ukeireyear = ukeiretime.strftime('%Y')
ukeireyearstrtoint = int(ukeireyear)
ukeireyeartoreiwa = ukeireyearstrtoint - 2018
ukeirereiwa = str(ukeireyeartoreiwa)
ukeiremonth = ukeiretime.strftime('%m').lstrip("0")
ukeirebi = ukeiretime.strftime('%d').lstrip("0")
bat01 = 'mkdir '+ 'R'+ ukeirereiwa +'.'+ ukeiremonth +'.'+ ukeirebi +'_'+genba
bat02 = 'cd ' + 'R'+ ukeirereiwa +'.'+ ukeiremonth +'.'+ ukeirebi +'_'+genba
bat03 = 'mkdir '+ zerozeroone
bat04 = 'cd ' +  zerozeroone
bat05 = 'mkdir 1.550'
bat06 = 'cd..'

csvfilename = ukeireday + genba + '.csv'
df= pd.read_excel("JISA1481-1分析データシート作業改善.xlsx",sheet_name='CSV',index_col=[0,1])
df.to_csv(csvfilename,encoding="utf-16")

with open('makefolders.bat', 'w') as f:
    print(bat01, file=f)

with open('makefolders.bat', 'a') as f:
    print(bat02, file=f)

with open('makefolders.bat', 'a') as f:
    print(bat03, file=f)

with open('makefolders.bat', 'a') as f:
    print(bat04, file=f)

with open('makefolders.bat', 'a') as f:
    print(bat05, file=f)

with open('makefolders.bat', 'a') as f:
    print(bat06, file=f)
        
j = 10
while sheet0.cell(row=j,column=5).value != None:
    bat07 = 'xcopy ' + zerozeroone + ' ' + sheet0.cell(row=j,column=4).value + '_"' + sheet0.cell(row=j,column=5).value + '"' + r' /E /I'
    with open('makefolders.bat', 'a') as f:
        print(bat07, file=f)
    j = j+1

command1 = 'makefolders.bat'
subprocess.call(command1)

command2 = 'del '+'makefolders.bat'
subprocess.run(command2, shell=True)
"""
