from docx.api import Document
import openpyxl
import docx
from docx import document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import tkinter as tk
from tkinter.constants import PAGES
import docx2pdf
from docx2pdf import convert
import os

wb_db = openpyxl.load_workbook('seigikou_db.xlsx')
sheet1 = wb_db['1_event_table']
sheet3 = wb_db['3_kouen_table']
sheet5 = wb_db['5_CPD_table']
sheet7 = wb_db['7_rireki_table']
sheet9 = wb_db['9_member_table']

#↓前回分取得なのでstr(j-1)にしています。
maxRow = sheet1.max_row +1
for j in reversed(range(1,maxRow)):
    if sheet1.cell(row=j, column=1).value != None:
        datecell = "b"+str(j-1)
        namecell = "f"+str(j-1)
        placecell = "e"+str(j-1)
        contentscell = "e"+str(j-1)
        startcell = "c"+str(j-1)
        endcell = "d"+str(j-1)
        onlinecell = "m"+str(j-1)
        teachercell = "d"+str(j-1)

        break

date_value_b = sheet1[datecell].value
date_value = date_value_b.strftime('%Y%m%d')
yearhen = date_value_b.strftime('%Y')
monthstr = date_value_b.strftime('%m').lstrip("0")
daystr = date_value_b.strftime('%d').lstrip("0")
datevalue_line = yearhen + '年' + monthstr + '月' + daystr + '日'

name_value = sheet1[namecell].value
place_value = sheet1[placecell].value
start_value = sheet1[startcell].value
end_value = sheet1[endcell].value
online_value = sheet1[onlinecell].value

contents_line = "演題：「"+sheet3[contentscell].value+"」"
teacher_line = "講師："+sheet3[teachercell].value+"　氏"
contents_value = contents_line+"\r"+teacher_line

#西暦表記にする場合は、上記のdatevalue_lineをコメントアウトから復活し、もう片方のブロックにあるdatevalue_lineをコメントアウトする。

#yearstrtoint = int(yearhen)
#yeartoreiwa = yearstrtoint - 2018
#reiwastr = str(yeartoreiwa)
#datevalue_line = '令和' + reiwastr + '年' + monthstr + '月' + daystr + '日'
#元号表記にする場合は、上記のブロックをコメントアウトから復活し、もう片方のdatevalue_lineをコメントアウトする。

root = tk.Tk()
root.title(u"枚数入力")
root.geometry("250x150")

label1 = tk.Label(root,text="枚数入力")
label1.place(x=20,y=20)
entry1 = tk.Entry(root,width="6")
entry1.place(x=120,y=20)
label1mai = tk.Label(root,text="枚")
label1mai.place(x=180,y=20)

var = tk.IntVar()
var.set(1)

label2 = tk.Label(root,text="オンライン開催")
label2.place(x=20,y=50)
radio2y = tk.Radiobutton(
    text="する",
    value=1,
    variable=var
)
radio2y.place(x=20,y=70)
radio2n = tk.Radiobutton(
    text="しない",
    value=0,
    variable=var
)
radio2n.place(x=80,y=70)

button3=tk.Button(root,width=10,text="作成実行")
button3.place(x=90,y=110)

#pagezで、entry1.get()中の値を保持。オンライン開催する場合はvar.get()が1

def click():
    print901 = var.get()
    pagez = entry1.get()
    pageint = int(pagez)
    i=1
    for i in range(pageint):
        doc0 = docx.Document("Newcpdformat2.docx")
        tbl0 = doc0.tables[0]
        tbl1 = doc0.tables[1]
        row0 = tbl0.rows[0]
        row1 = tbl1.rows[1]
        row4 = tbl1.rows[4]
        row5 = tbl1.rows[5]
        row6 = tbl1.rows[6]
        row8 = tbl1.rows[8]
        row9 = tbl1.rows[9]

        stri = str(i+1)
        savedwordname = '.\pdfsmerge\\'+stri.zfill(3)+'.docx'
        p0text ="No. YO"+stri.zfill(3)+"\r"+datevalue_line
        p0=row0.cells[0].add_paragraph(p0text) 
        p0.alignment=WD_ALIGN_PARAGRAPH.RIGHT

        row4.cells[1].text = datevalue_line
        row4.cells[5].text = str(start_value.strftime('%H:%M'))
        row5.cells[1].text = datevalue_line
        row5.cells[5].text = str(end_value.strftime('%H:%M'))
        row6.cells[1].text = name_value
        row8.cells[1].text = place_value
        row9.cells[1].text = contents_value

        doc0.save(savedwordname)   
    import gatherdocx2pdf

    pagez = entry1.get()
    pageint = int(pagez)
    ii=1
    for ii in range(pageint):
        strii = str(ii+1)
        os.remove('.\pdfsmerge\\'+strii.zfill(3)+'.docx')
        os.remove('.\pdfsmerge\\'+strii.zfill(3)+'.pdf')        

    if var.get() == 1:
        doc0 = docx.Document("Newcpdformat2.docx")
        tbl0 = doc0.tables[0]
        tbl1 = doc0.tables[1]
        row0 = tbl0.rows[0]
        row1 = tbl1.rows[1]
        row4 = tbl1.rows[4]
        row5 = tbl1.rows[5]
        row6 = tbl1.rows[6]
        row8 = tbl1.rows[8]
        row9 = tbl1.rows[9]
    
        savedwordname = '.\pdfsmerge\\901.docx'
        p0text ="No. "+online_value+"\r"+datevalue_line
        p0=row0.cells[0].add_paragraph(p0text) 
        p0.alignment=WD_ALIGN_PARAGRAPH.RIGHT

        row4.cells[1].text = datevalue_line
        row4.cells[5].text = str(start_value.strftime('%H:%M'))
        row5.cells[1].text = datevalue_line
        row5.cells[5].text = str(end_value.strftime('%H:%M'))
        row6.cells[1].text = name_value
        row8.cells[1].text = place_value
        row9.cells[1].text = contents_value

        doc0.save(savedwordname)
        convert(".\pdfsmerge\\901.docx",".\pdfsmerge\\online.pdf")
        os.remove('.\pdfsmerge\\901.docx')


    root.destroy()

button3["command"] = click

root.mainloop()